import os
import uuid

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from docx import Document
from docx.shared import Inches

def transclusionTXT(templatefilepath, resultDocumentDirectoryPath, dataEntities):
    if os.path.isfile(templatefilepath) == False:
        return false
    TransclusionId = uuid.uuid1();
    TransclusionDocumentNumber = 1;
    for dataEntity in dataEntities:
        resultDocumentPath = "{0}\\{1}_{2}.txt".format(resultDocumentDirectoryPath, TransclusionId, TransclusionDocumentNumber)
        with open(templatefilepath,"r") as template_file, open(resultDocumentPath,"w") as result_file:
            result_file.write(template_file.read().format(**dataEntity));
        TransclusionDocumentNumber+=1
    return True

def transclusionPDF(templatefilepath, resultDocumentDirectoryPath, dataEntities):
    if not os.path.isfile(templatefilepath):
        return False
    TransclusionId = uuid.uuid1();
    TransclusionDocumentNumber = 1;
    for dataEntity in dataEntities:
        resultDocumentPath = "{0}\\{1}_{2}.pdf".format(resultDocumentDirectoryPath, TransclusionId, TransclusionDocumentNumber)   
        doc = SimpleDocTemplate(resultDocumentPath, pagesize=letter)
        parts = []
        style = ParagraphStyle(
            name='Normal',
            fontName='Helvetica-Bold',
            fontSize=9,
        )
        with open(templatefilepath,"r") as template_file:
            for line in template_file:
                content = line.format(**dataEntity)
                parts.append(Paragraph(content, style))
        doc.build(parts)
        TransclusionDocumentNumber+=1
    return True

def transclusionDOCX(templatefilepath, resultDocumentDirectoryPath, dataEntities):
    if not os.path.isfile(templatefilepath):
        return False
    TransclusionId = uuid.uuid1();
    TransclusionDocumentNumber = 1;
    for dataEntity in dataEntities:
        resultDocumentPath = "{0}\\{1}_{2}.docx".format(resultDocumentDirectoryPath, TransclusionId, TransclusionDocumentNumber)   
        document = Document()
        document.add_heading(dataEntity["id"], 0)
        with open(templatefilepath,"r") as template_file:
            for line in template_file:
                document.add_paragraph(line.format(**dataEntity))
        document.save(resultDocumentPath)
        TransclusionDocumentNumber+=1
    return True